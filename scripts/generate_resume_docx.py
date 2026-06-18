from copy import deepcopy
import html
import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/jegalminhyuk/Downloads/경력기술서.docx")
OUTPUT = ROOT / "public" / "files" / "resume_제갈민혁.docx"
PDF_OUTPUT = ROOT / "public" / "files" / "resume_제갈민혁.pdf"
FONT_PATH = Path("/System/Library/Fonts/Supplemental/AppleGothic.ttf")
BUNDLED_NODE = Path(
    "/Users/jegalminhyuk/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"
)


def load_resume_items():
    source = (ROOT / "src" / "data" / "projects.ts").read_text(encoding="utf-8")
    marker = "export const resumeItems = "
    start = source.index(marker) + len(marker)
    bracket_start = source.index("[", start)
    depth = 0
    quote = None
    escaped = False
    end = bracket_start

    for index, char in enumerate(source[bracket_start:], start=bracket_start):
        if quote:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == quote:
                quote = None
            continue
        if char in ("'", '"', "`"):
            quote = char
            continue
        if char == "[":
            depth += 1
        elif char == "]":
            depth -= 1
            if depth == 0:
                end = index + 1
                break

    array_literal = source[bracket_start:end]
    script = f"const resumeItems = {array_literal}; console.log(JSON.stringify(resumeItems));"
    node = str(BUNDLED_NODE if BUNDLED_NODE.exists() else "node")
    result = subprocess.run(
        [node, "-e", script],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_run(paragraph, text, *, bold=False, size=9):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    return run


def set_lines(cell, lines, *, size=9, bold_first=False):
    clear_cell(cell)
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        add_run(paragraph, line, bold=bold_first and index == 0, size=size)


def tech_line(item):
    return "  ".join(f"• {tech.strip()}" for tech in item.get("tech", item["environment"].split(",")))


def add_label_body_lines(cell, label, text, *, size=8):
    label_paragraph = cell.add_paragraph()
    label_paragraph.paragraph_format.space_before = Pt(4)
    label_paragraph.paragraph_format.space_after = Pt(0)
    label_paragraph.paragraph_format.line_spacing = 1.15
    add_run(label_paragraph, label, bold=True, size=size)

    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.left_indent = Inches(0.12)
    body_paragraph.paragraph_format.space_after = Pt(2)
    body_paragraph.paragraph_format.line_spacing = 1.15
    add_run(body_paragraph, text, size=size)


def fill_table(table, item):
    header = f"{item['company']} / {item['role']} ({item['period']})"
    set_lines(table.rows[0].cells[0], [header], size=10, bold_first=True)
    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    role_lines = [
        "Frontend",
        "Developer",
    ]
    set_lines(table.rows[1].cells[0], role_lines, size=8, bold_first=True)
    set_lines(table.rows[1].cells[1], [f"{item['project']} / {item['projectType']}"], size=10, bold_first=True)

    clear_cell(table.rows[2].cells[1])
    add_label_body_lines(table.rows[2].cells[1], "개발환경", item["environment"])
    if item.get("description"):
        add_label_body_lines(table.rows[2].cells[1], "서비스 설명", item["description"])
    add_label_body_lines(table.rows[2].cells[1], "사용 기술", tech_line(item))
    add_label_body_lines(table.rows[2].cells[1], "Situation", item["situation"])
    add_label_body_lines(table.rows[2].cells[1], "Task", item["task"])
    add_label_body_lines(table.rows[2].cells[1], "Action", item["action"])
    add_label_body_lines(table.rows[2].cells[1], "Result", item["result"])


def remove_table(table):
    element = table._element
    element.getparent().remove(element)


def configure_document_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(9)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 16, "111827"),
        ("Heading 2", 12, "1f2937"),
        ("Heading 3", 10, "374151"),
    ]:
        style = document.styles[style_name]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_labeled_paragraph(document, label, text, *, indent_body=True):
    label_paragraph = document.add_paragraph()
    label_paragraph.paragraph_format.line_spacing = 1.2
    label_paragraph.paragraph_format.space_after = Pt(0)
    label_run = label_paragraph.add_run(label)
    label_run.bold = True
    label_run.font.name = "맑은 고딕"
    label_run.font.size = Pt(9)

    body_paragraph = document.add_paragraph()
    if indent_body:
        body_paragraph.paragraph_format.left_indent = Inches(0.16)
    body_paragraph.paragraph_format.line_spacing = 1.2
    body_paragraph.paragraph_format.space_after = Pt(5)
    body_run = body_paragraph.add_run(text)
    body_run.font.name = "맑은 고딕"
    body_run.font.size = Pt(9)


def build_docx_from_scratch(items):
    document = Document()
    configure_document_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("프로젝트 상세 경력기술서")
    title_run.bold = True
    title_run.font.name = "맑은 고딕"
    title_run.font.size = Pt(20)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    meta_run = meta.add_run("제갈민혁 / 프론트엔드 개발자 / 2024.09 - 2026.06")
    meta_run.font.name = "맑은 고딕"
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor.from_string("4b5563")

    document.add_heading("경력 요약", level=1)
    for line in [
        "(주)방배동밸리 소속으로 사용자 온보딩, 구매자·판매자 서비스, 관리자 백오피스, 예약 전환 플로우, 고객센터·마이페이지 개선을 담당해온 프론트엔드 개발자입니다.",
        "Next.js, React, TypeScript 기반의 서비스 화면과 운영 도구를 구현하며 상태 관리, API 연동, 입력 검증, 예외 처리, 공통 훅·유틸 구조화를 중심으로 기능 안정성과 재사용성을 높여왔습니다.",
        "지도 기반 탐색, 비회원 주문 진입, 엑셀 다운로드, 사이드바, 모달, 탭 구조 등 반복되는 사용자·운영자 흐름을 공통화해 개발 효율과 서비스 사용성을 함께 개선했습니다.",
    ]:
        document.add_paragraph(line)

    document.add_heading("프로젝트 상세", level=1)
    for item in items:
        document.add_heading(f"{item['project']} / {item['projectType']}", level=2)
        add_labeled_paragraph(document, "기간/역할", f"{item['period']} | {item['role']}")
        add_labeled_paragraph(document, "개발환경", item["environment"])
        add_labeled_paragraph(document, "사용 기술", tech_line(item))
        if item.get("description"):
            add_labeled_paragraph(document, "서비스 설명", item["description"])
        add_labeled_paragraph(document, "Situation", item["situation"])
        add_labeled_paragraph(document, "Task", item["task"])
        add_labeled_paragraph(document, "Action", item["action"])
        add_labeled_paragraph(document, "Result", item["result"])

    document.save(OUTPUT)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    items = load_resume_items()
    if TEMPLATE.exists():
        document = Document(TEMPLATE)

        title_table = document.tables[0]
        set_lines(title_table.rows[0].cells[0], ["프로젝트 상세 경력기술서", "제갈민혁"], size=18, bold_first=True)
        for paragraph in title_table.rows[0].cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        template_table = deepcopy(document.tables[1]._tbl)
        for table in list(document.tables[1:]):
            remove_table(table)

        body = document._body._element
        for item in items:
            cloned = deepcopy(template_table)
            body.append(cloned)
            fill_table(document.tables[-1], item)
            document.add_paragraph()

        document.save(OUTPUT)
    else:
        build_docx_from_scratch(items)
    build_pdf(items)
    print(OUTPUT)


def para(text, style):
    return Paragraph(html.escape(text).replace("\n", "<br/>"), style)


def rich_para(text, style):
    return Paragraph(text.replace("\n", "<br/>"), style)


def build_pdf(items):
    pdfmetrics.registerFont(TTFont("AppleGothic", str(FONT_PATH)))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontName="AppleGothic",
        fontSize=24,
        leading=30,
        alignment=1,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "ResumeMeta",
        parent=styles["Normal"],
        fontName="AppleGothic",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#4b5563"),
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName="AppleGothic",
        fontSize=8.4,
        leading=12,
    )
    body_indent_style = ParagraphStyle(
        "ResumeBodyIndent",
        parent=body_style,
        leftIndent=8,
        spaceAfter=4,
    )
    header_style = ParagraphStyle(
        "ResumeHeader",
        parent=body_style,
        fontSize=10,
        leading=13,
        textColor=colors.white,
    )
    role_style = ParagraphStyle(
        "ResumeRole",
        parent=body_style,
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#111827"),
    )

    doc = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="제갈민혁 프로젝트 상세 경력기술서",
        author="제갈민혁",
    )

    story = [
        para("프로젝트 상세 경력기술서", title_style),
        para("제갈민혁 / 프론트엔드 개발자 / 2024.09 ~ 2026.06", meta_style),
        Spacer(1, 10),
    ]

    for item in items:
        detail_flow = [rich_para(f"<b>{html.escape(item['project'])} / {html.escape(item['projectType'])}</b>", body_style)]
        if item.get("description"):
            detail_flow.extend(
                [
                    rich_para("<b>서비스 설명</b>", body_style),
                    para(item["description"], body_indent_style),
                ]
            )
        for label, value in [
            ("개발환경", item["environment"]),
            ("사용 기술", tech_line(item)),
            ("Situation", item["situation"]),
            ("Task", item["task"]),
            ("Action", item["action"]),
            ("Result", item["result"]),
        ]:
            detail_flow.extend([rich_para(f"<b>{label}</b>", body_style), para(value, body_indent_style)])

        table = Table(
            [
                [para(f"{item['company']} / {item['role']} ({item['period']})", header_style)],
                [
                    rich_para(
                        "Frontend<br/>Developer",
                        role_style,
                    ),
                    detail_flow,
                ],
            ],
            colWidths=[42 * mm, 114 * mm],
            repeatRows=0,
        )
        table.setStyle(
            TableStyle(
                [
                    ("SPAN", (0, 0), (-1, 0)),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
                    ("BACKGROUND", (0, 1), (0, 1), colors.HexColor("#f3f4f6")),
                    ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#111827")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d1d5db")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]
            )
        )
        story.append(KeepTogether([table, Spacer(1, 9)]))

    doc.build(story)


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"failed: {exc}", file=sys.stderr)
        raise
